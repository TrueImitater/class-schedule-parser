import os

import argparse
import docx
from docx.shared import Cm

from docx_parser import parse_table
from parser_utils import TableCoreParams

from table_utils import (set_page_size_a3,
                         create_table_head,
                         fill_schedule_table,
                         set_column_width)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()

    parser.add_argument("-i", "--input-file", type=str,
                        help="Путь до файла с расписанием", required=True)
    parser.add_argument("-g", "--groups", nargs="+", type=str,
                        help="Наименование групп, расписание которых необходимо распарсить")
    parser.add_argument("-o", "--output-dir", default=r".", type=str,
                        help="Директория, в которую будут сохранены документы с расписаниями")
    parser.add_argument("-all", "--parse-all", action='store_true',
                        help="Флаг, при установке которого происходит парсинг расписаний всех найденных в загруженном файле групп")

    # парсим параметры из командной строки
    input_params = parser.parse_args()

    # пробуем прочитать файл
    docx_file = docx.Document(input_params.input_file)

    # создаем объект, содержащий основные параметры по анализируемой таблице
    table_core_params = TableCoreParams(docx_file.tables[0])

    # получаем массив групп, расписание которых необходимо распарсить
    target_groups = table_core_params.group_name_indexes.keys(
    ) if input_params.parse_all else input_params.groups

    # если группы не указаны
    if target_groups is None:
        raise Exception(
            "[ERROR] Не были указаны группы, расписание которых необходимо распарсить")

    # для каждой из указанной пользователем группы создаем таблицу
    for group_name in target_groups:
        # если группа присутствует в загруженном файле
        if table_core_params.group_name_indexes.get(group_name, None) is not None:
            schedule_dict = parse_table(table_core_params.table.rows[table_core_params.start_row_idx:], table_core_params.day_col_idx,
                                        table_core_params.time_col_idx, table_core_params.group_name_indexes[group_name])

            doc = docx.Document()  # создаем документ

            # изменяем ширину полей документа
            section = doc.sections[0]
            section.top_margin = Cm(1.5)
            section.left_margin = Cm(2)

            # устанавливаем размер страницы A3
            set_page_size_a3(section)

            # создаем шапку таблицы
            create_table_head(doc, group_name)
            fill_schedule_table(doc.tables[0], schedule_dict)

            # устанавливаем ширину столбцов полученной таблицы
            set_column_width(doc.tables[0], 0, 3.5)
            set_column_width(doc.tables[0], 1, 2.5)
            set_column_width(doc.tables[0], 2, 8.5)
            set_column_width(doc.tables[0], 3, 8.5)

            # сохраняем документ
            try:
                doc.save(os.path.join(
                    input_params.output_dir, f"{group_name}.docx"))
            except PermissionError:
                print(
                    f"[ERROR] Не удалось сохранить файл {group_name}.docx по пути {input_params.output_dir}. Возможно, он уже открыт в другой программе")
