""" Функции для оформления таблиц """

import docx

import docx.enum
import docx.enum.table
import docx.enum.text
import docx.oxml
import docx.oxml.shape
import docx.oxml.shared
import docx.shared
import docx.table
from docx.shared import Inches, Cm

from docx import Document
from parser_utils import *

TEXT_FONT_KEY_WORD = "Times New Roman"

DEFAULT_CELL_COLOR = "#e1ffe1"

UPPER_HEAD_CELL_COLOR = "#ff9999"
UPPER_CELL_COLOR = "#ffcccc"

LOWER_HEAD_CELL_COLOR = "#99ccff"
LOWER_CELL_COLOR = "#ccecff"


def change_cell_style(cur_cell: docx.table._Cell, font_name: str,
                      cell_size: int):
    """ Изменение общих характеристик ячейки """
    cur_cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    cur_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    cur_cell_font = cur_cell.paragraphs[0].runs[0].font
    cur_cell_font.name = font_name
    cur_cell_font.size = docx.shared.Pt(cell_size)


def set_cell_background_color(cell, color):
    """ Функция для изменения цвета фона ячейки """
    # Получаем доступ к XML ячейки
    cell_properties = cell._element.get_or_add_tcPr()
    # Создаем элемент для заливки цвета
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    # Добавляем элемент к свойствам ячейки
    cell_properties.append(shading)


def create_table_head(doc: docx.Document):
    """ Создание шапки таблицы с распершеннным расписанием """

    # создаем экземпляр таблицы
    schedule_table = doc.add_table(rows=2, cols=4)
    schedule_table.style = 'Table Grid'

    schedule_table.cell(0, 0).text = "Группа"
    schedule_table.cell(0, 0).merge(schedule_table.cell(0, 1))

    schedule_table.cell(1, 0).text = "день"
    schedule_table.cell(1, 1).text = "время"

    schedule_table.cell(0, 2).text = "Верхняя"
    schedule_table.cell(0, 2).merge(schedule_table.cell(1, 2))
    schedule_table.cell(0, 3).text = "Нижняя"
    schedule_table.cell(0, 3).merge(schedule_table.cell(1, 3))

    # оформляем ячейку "Группа"
    group_cell = schedule_table.cell(0, 0)
    group_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    change_cell_style(group_cell, TEXT_FONT_KEY_WORD, 12)
    set_cell_background_color(group_cell, DEFAULT_CELL_COLOR)

    # оформляем ячейку "День"
    day_of_week_cell = schedule_table.cell(1, 0)
    day_of_week_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    change_cell_style(day_of_week_cell, TEXT_FONT_KEY_WORD, 12)
    set_cell_background_color(day_of_week_cell, DEFAULT_CELL_COLOR)

    # оформляем ячейку "Время"
    time_cell = schedule_table.cell(1, 1)
    time_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    change_cell_style(time_cell, TEXT_FONT_KEY_WORD, 12)
    set_cell_background_color(time_cell, DEFAULT_CELL_COLOR)

    # оформляем ячейку "Верхняя"
    upper_cell = schedule_table.cell(0, 2)
    upper_cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    upper_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    upper_cell_font = upper_cell.paragraphs[0].runs[0].font
    upper_cell_font.name = TEXT_FONT_KEY_WORD
    upper_cell_font.size = docx.shared.Pt(16)
    upper_cell_font.bold = True
    set_cell_background_color(upper_cell, UPPER_HEAD_CELL_COLOR)

    # оформляем ячейку "Нижняя"
    lower_cell = schedule_table.cell(0, 3)
    lower_cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    lower_cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    lower_cell_font = lower_cell.paragraphs[0].runs[0].font
    lower_cell_font.name = TEXT_FONT_KEY_WORD
    lower_cell_font.size = docx.shared.Pt(16)
    lower_cell_font.bold = True
    set_cell_background_color(lower_cell, LOWER_HEAD_CELL_COLOR)


def add_schedule_row(schedule_table: docx.table.Table, time: str, upper_schedule: str, lower_schedule: str):
    """ Добавление строки в таблицу """

    cells = schedule_table.add_row().cells

    # # оформляем ячейку "День"
    # change_cell_style(cells[0], TEXT_FONT_KEY_WORD, 10)
    set_cell_background_color(cells[0], DEFAULT_CELL_COLOR)

    # оформляем ячейку "Время"
    cells[1].text = time
    change_cell_style(cells[1], TEXT_FONT_KEY_WORD, 11)
    set_cell_background_color(cells[1], DEFAULT_CELL_COLOR)

    # оформляем ячейку "Верхней недели"
    cells[2].text = upper_schedule
    change_cell_style(cells[2], TEXT_FONT_KEY_WORD, 11)
    set_cell_background_color(
        cells[2], DEFAULT_CELL_COLOR if upper_schedule == "" else UPPER_CELL_COLOR)

    # оформляем ячейку "Нижней недели"
    cells[3].text = lower_schedule
    change_cell_style(cells[3], TEXT_FONT_KEY_WORD, 11)
    set_cell_background_color(
        cells[3], DEFAULT_CELL_COLOR if lower_schedule == "" else LOWER_CELL_COLOR)


def fill_schedule_table(schedule_table: docx.table.Table, schedule_dict: dict):
    """ Заполнение таблицы информацией о парах """
    start_row_idx = cur_row_idx = 2

    for (key, value) in schedule_dict.items():
        cur_day = key.capitalize()

        cur_sub_idx = 0

        if (len(value) == 1):
            continue

        while True:
            try:
                first_time = value[cur_sub_idx][0]
                first_schedule = value[cur_sub_idx][1]
            except:
                print(cur_sub_idx, value)
                return
            if first_time == END_OF_DAY:
                break

            second_time = value[cur_sub_idx + 1][0]
            second_schedule = value[cur_sub_idx + 1][1]

            if (first_time == second_time):
                # если это верхняя и нижняя неделя
                if (first_schedule == second_schedule == ""):
                    # если пар нет
                    pass
                else:
                    add_schedule_row(schedule_table, first_time,
                                     first_schedule, second_schedule)
                    cur_row_idx += 1
                cur_sub_idx += 2
            else:
                # если это одновременно и верхняя и нижняя недели
                if (first_schedule != ""):
                    # если это пара по обеим неделям - занести ее в таблицу
                    add_schedule_row(schedule_table, first_time,
                                     first_schedule, first_schedule)
                    cur_row_idx += 1
                else:
                    pass
                cur_sub_idx += 1

        if (cur_row_idx == start_row_idx):
            add_schedule_row(schedule_table, "", "", "")
            schedule_table.rows[-1].cells[0].text = cur_day

            change_cell_style(
                schedule_table.rows[-1].cells[0], TEXT_FONT_KEY_WORD, 11)

            start_row_idx += 1
            cur_row_idx += 1
        else:
            schedule_table.rows[cur_row_idx - 1].cells[0].text = cur_day
            change_cell_style(
                schedule_table.rows[-1].cells[0], TEXT_FONT_KEY_WORD, 11)
            schedule_table.cell(start_row_idx, 0).merge(
                schedule_table.cell(cur_row_idx - 1, 0))

            schedule_table.rows[-1].cells[0].text = cur_day
            schedule_table.cell(
                start_row_idx, 0).vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
            schedule_table.cell(
                start_row_idx, 0).paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            start_row_idx = cur_row_idx


def set_page_size_a3(section):
    """ Установка размера страницы A3 """
    section.page_width = Cm(29.7)
    section.page_height = Cm(42)


def set_cell_border(cell: docx.table._Cell, side: str, sz: int = 4, color: str = "000000"):
    """Установка границы ячейки"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Получаем существующие границы или создаем новые
    tcBorders = tcPr.xpath('./w:tcBorders')

    print(
        f"{'1' if tcBorders else '0'} Горизонтальный span: {tc.grid_span}\tВертикальный span: {tc.vMerge}")

    # if tcBorders:
    #     return

    if not tcBorders:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    else:
        tcBorders = tcBorders[0]

    # Создаем и добавляем границу
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:sz'), str(sz))
    border.set(qn('w:val'), 'single')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def make_bold_line(table: Table, index: int, side: str, is_row: bool = True, sz: int = 16):
    """
    Делает жирной строку или столбец таблицы
    Args:
        table: Таблица
        index: Индекс строки/столбца
        is_row: True - строка, False - столбец
        sz: Толщина линии
    """

    iterable_cells = table.rows[index].cells if is_row else table.columns[index].cells

    print("--- Старт ---")

    count: int = 0

    for cell in iterable_cells:
        set_cell_border(cell, side, sz)

        count += 1

    print(f"--- Конец --- count = {count}")


def set_table_bold_borders(table: Table):
    """
    Устанавливает жирные границы таблицы
    """

    # # Для начала все ребра столбцов должны быть жирными
    # for i in range(len(table.columns)):
    #     make_bold_line(table, i, "left", False)
    #     make_bold_line(table, i, "right", False)

    # # первые две строки - жирные и низ последней
    # make_bold_line(table, 0, "top")
    # make_bold_line(table, 1, "top")
    # make_bold_line(table, 1, "bottom")
    # make_bold_line(table, len(table.rows) - 1, "bottom")


def set_column_width(table: docx.table.Table, col_idx: int, width: float):
    """
    Установка ширины колонки
    Args:
        table: Таблица
        col_idx: Индекс колонки
        width: Ширина в сантиметрах
    """
    for row in table.rows:
        row.cells[col_idx].width = Cm(width)


def add_doc_space(doc, top_cm: float = 0, bottom_cm: float = 0, left_cm: float = 0, right_cm: float = 0):
    """Добавление отступов по всем сторонам"""
    section = doc.sections[-1]
    section.top_margin += Cm(top_cm)
    section.bottom_margin += Cm(bottom_cm)
    section.left_margin += Cm(left_cm)
    section.right_margin += Cm(right_cm)


def create_table(schedule_dict: dict, save_path: str) -> bool:
    """ Создание таблицы по словарю с парами по дням недели """
    doc = Document()

    # изменяем ширину полей документа
    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(0.5)

    # устанавливаем размер страницы A3
    set_page_size_a3(section)

    # создаем шапку таблицы
    create_table_head(doc)

    fill_schedule_table(doc.tables[0], schedule_dict)
    set_table_bold_borders(doc.tables[0])

    # устанавливаем ширину столбцов полученной таблицы
    set_column_width(doc.tables[0], 0, 3.5)
    set_column_width(doc.tables[0], 1, 2.5)
    set_column_width(doc.tables[0], 2, 8.5)
    set_column_width(doc.tables[0], 3, 8.5)

    # измнение отступы таблицы (через изменение ширины поля документа)
    add_doc_space(doc, top_cm=1, left_cm=2.5)

    doc.save(save_path)

    return True
